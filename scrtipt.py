import docx
import os
import json
import re

paths = []
folder = os.getcwd()
for root, dirs, files in os.walk(folder):
    for file in files:
        if file.endswith('docx') and not file.startswith('~'):
            paths.append(os.path.join(root, file))

f1 = open('output_1.txt', 'w')
doc = docx.Document(paths[0])

pattern = re.compile('w:fill=\"(\S*)\"')
for table in doc.tables:
    len_rows = len(table.rows)
    is_in_apart = False
    apartname = ''
    apartspec = ''
    for row in table.rows:
        len_cells = len(row.cells)
        if len_cells == 5:
            cell_0 = row.cells[0]
            match = pattern.search(cell_0._tc.xml)
            if match:
                # print(match.group(1))
                if match.group(1) == 'auto' or match.group(1) == 'ffffff':
                    if (cell_0.text != apartspec):
                        apartspec = cell_0.text
                        f1.write('  apart_spec: ' + row.cells[0].text + '\n')
                    # print('    apart_floor: ' + row.cells[1].text + '\n    apart_price: ' + row.cells[2].text + '\n')
                    f1.write('    apart_floor: ' + row.cells[1].text + '\n    apart_price: ' + row.cells[4].text + '\n')
                else:
                    if (cell_0.text != apartname):
                        # print(cell_0.text + '=' + apartname)
                        apartname = cell_0.text
                        f1.write('\napart_name: ' + row.cells[0].text + '\n')
        if len_cells == 3:
            cell_0 = row.cells[0]
            match = pattern.search(cell_0._tc.xml)
            if match:
                # print(match.group(1))
                if match.group(1) == 'auto' or match.group(1) == 'ffffff':
                    if (cell_0.text != apartspec):
                        apartspec = cell_0.text
                        f1.write('  apart_spec: ' + row.cells[0].text + '\n')
                    f1.write('    apart_floor: ' + row.cells[1].text + '\n    apart_price: ' + row.cells[2].text + '\n')
                else:
                    if (cell_0.text != apartname):
                        # print(cell_0.text + '=' + apartname)
                        apartname = cell_0.text
                        f1.write('\napart_name: ' + row.cells[0].text + '\n')
        elif len_cells == 2:
            cell_0 = row.cells[0]
            match = pattern.search(cell_0._tc.xml)
            if match:
                if match.group(1) == 'auto':
                    if (cell_0.text != apartspec):
                        apartspec = cell_0.text
                        f1.write('apart_spec: ' + row.cells[0].text + '\n')
                    f1.write('   apart_price: ' + row.cells[1].text + '\n')
                        # for i in range(1, len_cells-1):
                        #     f1.write('   ' + row.cells[i].text + '   ')
                else:            
                    if (cell_0.text != apartname):
                        apartname = cell_0.text
                        f1.write('\napart_name: ' + apartname + '\n')
        f1.write('\n')
    f1.write('------------------\n')
                



# pattern = re.compile('w:fill=\"(\S*)\"')
# for table in doc.tables:
#     unique = set()
#     len_rows = len(table.rows)
#     level = 0
#     for row in table.rows:
#         len_cells = len(row.cells)
#         for cell in row.cells:
#             tc = cell._tc
#             cell_loc = (tc.top, tc.bottom, tc.left, tc.right)
#             if cell_loc in unique:
#                 f1.write(' ')
#             else:
#                 unique.add(cell_loc)
#                 match = pattern.search(cell._tc.xml)
#                 if match:
#                     if match.group(1) != 'auto':
#                         f1.write('name:   ' + cell.text + ' //\n')
#                     else:            
#                         f1.write(cell.text)
                    
#                     if cell != row.cells[len_cells-1]:
#                         f1.write(' | ')
#         f1.write('\n')
                



# f2 = open('output_2.txt', 'w')

# doc = docx.Document(paths[1])

# for table in doc.tables:
#     unique, merged = set(), set()
#     for row in table.rows:
#         for cell in row.cells:
#             match = pattern.search(cell._tc.xml)
#             if match:
#                 if match.group(1) != 'FFFFFF':
#                     f2.write(cell.text + ' // \n')
#             else:            
#                 f2.write(cell.text)
            
#             if cell != row.cells[-1]:
#                 f2.write(' | ')
#         f2.write('\n')

# f3 = open('output_3.txt', 'w')

# doc = docx.Document(paths[2])

# for table in doc.tables:
#     unique, merged = set(), set()
#     for row in table.rows:
#         for cell in row.cells:
#             match = pattern.search(cell._tc.xml)
#             if match:
#                 if match.group(1) != 'FFFFFF':
#                     f3.write(cell.text + ' // \n')
#             else:            
#                 f3.write(cell.text)
            
#             if cell != row.cells[-1]:
#                 f3.write(' | ')
#         f3.write('\n')